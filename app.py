import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import calendar
import openpyxl
from docx import Document
from docx.shared import Inches
import re
import io
import zipfile
import os
from pathlib import Path
import tempfile
import base64
import warnings
warnings.filterwarnings('ignore')

# Import configuration
import config

# Helper function for logo
def get_base64_of_image(path):
    """Convert image to base64 string for embedding in HTML"""
    try:
        with open(path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    except Exception as e:
        return ""

# Try to import docx2pdf for PDF generation
try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Configure Streamlit page
st.set_page_config(
    page_title="نظام إدارة الدورات التدريبية",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': None,
        'Report a bug': None,
        'About': None
    }
)

# RTL CSS styling with enhanced purple-teal theme
st.markdown("""
<style>
    /* Import Google Fonts for Arabic */
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@300;400;600;700&display=swap');
    
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden !important;}
    footer {visibility: hidden !important;}
    header {visibility: hidden !important;}
    .stDeployButton {display: none !important;}
    .stActionButton {display: none !important;}
    
    /* Hide any "Created by" or attribution text */
    .css-hi6a2p {display: none !important;}
    .css-18e3th9 {display: none !important;}
    .css-1d391kg .css-12oz5g7 {display: none !important;}
    [data-testid="stSidebarNavItems"] {display: none !important;}
    [data-testid="stSidebarNav"] {display: none !important;}
    
    /* Hide any footer or attribution elements */
    .css-1y4p8pa {display: none !important;}
    .css-1dp5vir {display: none !important;}
    
    /* Additional hiding for any remaining branding */
    .css-1rs6os {display: none !important;}
    .css-17eq0hr {display: none !important;}
    
    /* Mobile specific hiding */
    @media screen and (max-width: 768px) {
        #MainMenu {visibility: hidden !important;}
        footer {visibility: hidden !important;}
        header {visibility: hidden !important;}
        .stDeployButton {display: none !important;}
        .stActionButton {display: none !important;}
        
        /* Mobile specific selectors for "Created by" */
        .css-1y4p8pa {display: none !important;}
        .css-1dp5vir {display: none !important;}
        .css-hi6a2p {display: none !important;}
        .css-18e3th9 {display: none !important;}
        [data-testid="stSidebarNavItems"] {display: none !important;}
        [data-testid="stSidebarNav"] {display: none !important;}
        
        /* Hide mobile menu and navigation */
        .css-1lcbmhc {display: none !important;}
        .css-1y0tads {display: none !important;}
        .css-17ziqus {display: none !important;}
        .css-1544g2n {display: none !important;}
        
        /* Hide any mobile branding footer */
        .stApp > footer {display: none !important;}
        .main-footer {display: none !important;}
        .streamlit-footer {display: none !important;}
        
        /* Additional mobile-specific branding elements */
        .css-1rs6os {display: none !important;}
        .css-17eq0hr {display: none !important;}
        .css-12oz5g7 {display: none !important;}
        .css-1d391kg {display: none !important;}
        
        /* Hide any "Made with Streamlit" text on mobile */
        [data-testid="stBottomBlockContainer"] {display: none !important;}
        [data-testid="stFooter"] {display: none !important;}
        
        /* Target the specific mobile footer container from your screenshot */
        .css-1outpf7 {display: none !important;}
        .css-16huue1 {display: none !important;}
        .css-1v0mbdj {display: none !important;}
        .css-1wbqy5l {display: none !important;}
        
        /* Hide the entire bottom container on mobile */
        div[data-testid="stBottom"] {display: none !important;}
        div[data-testid="stBottomContainer"] {display: none !important;}
        
        /* Hide any element containing "Created by" text */
        div:contains("Created by") {display: none !important;}
        span:contains("Created by") {display: none !important;}
        p:contains("Created by") {display: none !important;}
        
        /* Hide streamlit cloud footer */
        .st-emotion-cache-1wbqy5l {display: none !important;}
        .st-emotion-cache-16huue1 {display: none !important;}
        .st-emotion-cache-1outpf7 {display: none !important;}
        
        /* Additional emotion cache classes for mobile */
        [class*="emotion-cache"] div:contains("Created") {display: none !important;}
        [class*="st-emotion"] div:contains("Created") {display: none !important;}
    }
    
    /* Main layout and RTL support */
    .main {
        direction: rtl;
        text-align: right;
        font-family: 'Cairo', sans-serif;
    }
    .stSelectbox > div > div {
        direction: rtl;
    }
    .stDataFrame {
        direction: rtl;
    }
    
    /* Header with gradient background */
    .header-container {
        background: linear-gradient(135deg, #6A3CBC 0%, #2EC4B6 100%);
        padding: 2rem 1rem;
        border-radius: 15px;
        margin-bottom: 2rem;
        box-shadow: 0 8px 32px rgba(106, 60, 188, 0.3);
        text-align: center;
        position: relative;
        overflow: hidden;
    }
    
    .header-container::before {
        content: '';
        position: absolute;
        top: -50%;
        right: -50%;
        width: 200%;
        height: 200%;
        background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%);
        animation: float 6s ease-in-out infinite;
    }
    
    @keyframes float {
        0%, 100% { transform: translate(0, 0) rotate(0deg); }
        50% { transform: translate(-20px, -10px) rotate(180deg); }
    }
    
    .header-title {
        color: white;
        font-size: 2.5rem;
        font-weight: 700;
        margin: 0.5rem 0;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        position: relative;
        z-index: 2;
    }
    
    .header-subtitle {
        color: rgba(255,255,255,0.9);
        font-size: 1.2rem;
        font-weight: 400;
        margin: 0;
        position: relative;
        z-index: 2;
    }
    
    .logo-container {
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 1rem;
        margin-bottom: 1rem;
        position: relative;
        z-index: 2;
    }
    
    .logo-image {
        width: 80px;
        height: 80px;
        border-radius: 15px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
        background: rgba(255,255,255,0.1);
        padding: 0.5rem;
    }
    
    /* Enhanced metric cards */
    .metric-card {
        background: linear-gradient(135deg, #2D2F47 0%, #3A3D5C 100%);
        border: 1px solid rgba(106, 60, 188, 0.3);
        padding: 1.5rem;
        border-radius: 15px;
        text-align: center;
        margin: 0.5rem 0;
        box-shadow: 0 4px 20px rgba(106, 60, 188, 0.2);
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .metric-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 3px;
        background: linear-gradient(90deg, #6A3CBC 0%, #2EC4B6 100%);
    }
    
    .metric-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 8px 30px rgba(106, 60, 188, 0.4);
    }
    
    .metric-value {
        font-size: 2.5rem;
        font-weight: 700;
        color: #2EC4B6;
        margin-bottom: 0.5rem;
        text-shadow: 0 2px 4px rgba(0,0,0,0.3);
    }
    
    .metric-label {
        font-size: 1rem;
        color: rgba(255,255,255,0.8);
        font-weight: 500;
    }
    
    /* Status indicators */
    .status-match {
        color: #2EC4B6;
        font-weight: bold;
    }
    .status-missing {
        color: #FF6B6B;
        font-weight: bold;
    }
    .status-unused {
        color: #FFB347;
        font-weight: bold;
    }
    
    /* Headers */
    h1, h2, h3 {
        text-align: right;
        direction: rtl;
        color: #FFFFFF;
        font-family: 'Cairo', sans-serif;
    }
    
    h1 {
        color: #6A3CBC;
        font-weight: 700;
    }
    
    h2 {
        color: #2EC4B6;
        font-weight: 600;
    }
    
    /* Buttons enhancement */
    .stButton > button {
        width: 100%;
        background: linear-gradient(135deg, #6A3CBC 0%, #2EC4B6 100%);
        border: none;
        border-radius: 10px;
        color: white;
        font-weight: 600;
        padding: 0.75rem 1.5rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(106, 60, 188, 0.3);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(106, 60, 188, 0.4);
    }
    
    /* Sidebar enhancement */
    .css-1d391kg {
        background: linear-gradient(180deg, #2D2F47 0%, #1E1E2F 100%);
    }
    
    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1rem;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: linear-gradient(135deg, #2D2F47 0%, #3A3D5C 100%);
        border-radius: 10px;
        color: #FFFFFF;
        border: 1px solid rgba(106, 60, 188, 0.3);
        padding: 0.75rem 1.5rem;
        font-weight: 600;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #6A3CBC 0%, #2EC4B6 100%);
        color: white;
        box-shadow: 0 4px 15px rgba(106, 60, 188, 0.3);
    }
    
    /* Data tables */
    .stDataFrame {
        border-radius: 10px;
        overflow: hidden;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
    }
    
    /* Success/Error messages */
    .stSuccess {
        background: linear-gradient(135deg, #2EC4B6 0%, #20A39E 100%);
        border: none;
        border-radius: 10px;
    }
    
    .stError {
        background: linear-gradient(135deg, #FF6B6B 0%, #E55353 100%);
        border: none;
        border-radius: 10px;
    }
    
    .stWarning {
        background: linear-gradient(135deg, #FFB347 0%, #FF9F00 100%);
        border: none;
        border-radius: 10px;
    }
    
    /* File uploader styling */
    .stFileUploader {
        border: 2px dashed rgba(106, 60, 188, 0.3);
        border-radius: 10px;
        padding: 1rem;
        background: rgba(45, 47, 71, 0.3);
    }
    
    /* Plotly charts enhancement */
    .js-plotly-plot {
        border-radius: 10px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
    }
</style>

<script>
// JavaScript to aggressively hide "Created by" elements on mobile
function hideCreatedByElements() {
    // Hide any element containing "Created by" text
    const allElements = document.querySelectorAll('*');
    allElements.forEach(element => {
        if (element.textContent && element.textContent.includes('Created by')) {
            element.style.display = 'none';
            element.style.visibility = 'hidden';
            element.remove();
        }
        if (element.textContent && element.textContent.includes('Hosted with Streamlit')) {
            element.style.display = 'none';
            element.style.visibility = 'hidden';
            element.remove();
        }
    });
    
    // Hide common footer selectors
    const footerSelectors = [
        '[data-testid="stBottom"]',
        '[data-testid="stFooter"]',
        '.css-1outpf7',
        '.css-16huue1',
        '.css-1wbqy5l',
        '.st-emotion-cache-1wbqy5l',
        '.st-emotion-cache-16huue1',
        'footer',
        '.streamlit-footer'
    ];
    
    footerSelectors.forEach(selector => {
        const elements = document.querySelectorAll(selector);
        elements.forEach(el => {
            el.style.display = 'none';
            el.style.visibility = 'hidden';
        });
    });
}

// Run immediately
hideCreatedByElements();

// Run when DOM changes (for dynamic content)
const observer = new MutationObserver(hideCreatedByElements);
observer.observe(document.body, { childList: true, subtree: true });

// Run on load and resize
window.addEventListener('load', hideCreatedByElements);
window.addEventListener('resize', hideCreatedByElements);

// Run every 500ms to catch any late-loading elements
setInterval(hideCreatedByElements, 500);
</script>
""", unsafe_allow_html=True)

# ========================= DATA LOADING FUNCTIONS =========================

@st.cache_data
def load_excel_data(file_path, month_sheet=None):
    """
    Load Excel data from specified file and sheet
    Returns DataFrame with course data
    """
    try:
        if month_sheet:
            df = pd.read_excel(file_path, sheet_name=month_sheet)
        else:
            # Load first sheet if no specific sheet mentioned
            df = pd.read_excel(file_path)
        
        # Clean the data - remove empty rows and unnecessary columns
        # Remove completely empty rows
        df = df.dropna(how='all')
        
        # Remove columns that are completely empty or just "Unnamed"
        df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
        
        # Filter rows that have actual course data (should have course name)
        course_name_cols = [col for col in df.columns if 'اسم الدورة' in str(col) and 'تاريخ' not in str(col)]
        if not course_name_cols:
            # Fallback: look for any column with 'اسم' and 'برنامج' or just 'اسم' and has substantial data
            course_name_cols = [col for col in df.columns if 'اسم' in str(col) and ('برنامج' in str(col) or 'دورة' in str(col)) and 'تاريخ' not in str(col)]
        
        if course_name_cols:
            course_col = course_name_cols[0]
            # Keep only rows that have course names
            df = df[df[course_col].notna() & (df[course_col] != '') & (df[course_col].astype(str).str.strip() != '')]
        else:
            # If no clear course column, just remove obviously empty rows
            # Keep rows where at least 3 columns have data
            df = df[df.count(axis=1) >= 3]
        
        # Ensure date columns are properly parsed
        date_columns = [col for col in df.columns if 'تاريخ' in str(col).lower() or 'date' in str(col).lower()]
        for col in date_columns:
            try:
                df[col] = pd.to_datetime(df[col], errors='coerce')
            except:
                pass
        
        # Reset index after filtering
        df = df.reset_index(drop=True)
        
        return df
    except Exception as e:
        st.error(f"خطأ في تحميل البيانات: {str(e)}")
        return pd.DataFrame()

def get_available_sheets(file_path):
    """
    Get list of available sheets in Excel file
    """
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True)
        sheets = wb.sheetnames
        wb.close()
        return sheets
    except Exception as e:
        st.error(f"خطأ في قراءة أوراق العمل: {str(e)}")
        return []

# ========================= WORD TEMPLATE FUNCTIONS =========================

def extract_placeholders_from_word(docx_path):
    """
    Extract all placeholders from Word document
    Returns list of placeholders found
    Safe version that handles table errors gracefully
    """
    try:
        doc = Document(docx_path)
        placeholders = set()
        
        # Extract from paragraphs (safe)
        for para in doc.paragraphs:
            try:
                matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
                placeholders.update(matches)
            except Exception as para_error:
                continue
        
        # Extract from tables with enhanced error handling
        try:
            for table_idx, table in enumerate(doc.tables):
                try:
                    for row_idx, row in enumerate(table.rows):
                        try:
                            for cell_idx, cell in enumerate(row.cells):
                                try:
                                    # Safely access cell text
                                    cell_text = cell.text
                                    matches = re.findall(r'\{\{([^}]+)\}\}', cell_text)
                                    placeholders.update(matches)
                                except Exception as cell_error:
                                    # Skip problematic cells (like grid_offset errors)
                                    continue
                        except Exception as row_error:
                            # Skip problematic rows
                            continue
                except Exception as table_error:
                    # Skip problematic tables
                    continue
        except Exception as tables_error:
            # If table processing fails completely, continue with paragraphs only
            st.warning("⚠️ تخطي معالجة الجداول بسبب خطأ في التنسيق")
        
        return list(placeholders)
    except Exception as e:
        st.error(f"خطأ في قراءة قالب Word: {str(e)}")
        return []

def build_mapping(row, df_columns):
    """
    Build mapping between Excel columns and Word Content Control tags
    Content Control Tags available:
    - اسم البرنامج, اسم الدورة, الفئة المستهدفة, طريقة الطرح
    - العدد المتوقع, وقت الدورة/البرنامج, تاريخ التنفيذ, مدتها
    - اســــم الــمــدرب, الـــمـتـطــلــبــات, تنفيذ البرنامج/الدورة, مقر التنفيذ
    """
    mapping = {}
    
    # Define Excel to Content Control tag mappings
    excel_to_tag_mapping = {
        # Excel Column Name -> Content Control Tag
        'اسم الدورة بالعربي': 'اسم الدورة',
        'الفئة المستهدفة': 'الفئة المستهدفة',
        'طريقة الطرح': 'طريقة الطرح',
        'اسم المدرب': 'اســــم الــمــدرب',
        'مكان الانعقاد ': 'مقر التنفيذ',  # Fixed: exact match with trailing space
        'مكان الانعقاد': 'مقر التنفيذ',   # Alternative without space
        'الوقت': 'وقت الدورة/البرنامج',
        'عدد الايام': 'مدتها',
        'تاريخ بداية الدورة بالميلادي': 'تاريخ التنفيذ',
        'تاريخ بداية الدورة بالهجري': 'تاريخ التنفيذ',
        'تحتاج لمعمل؟': 'استخدام معمل الحاسب',  # Fixed: lab field mapping
        
        # Additional mappings for other available fields
        'جهة التدريب': 'تنفيذ البرنامج/الدورة',
        'اسم جهة التدريب': 'تنفيذ البرنامج/الدورة',
    }
    
    # Build mapping from Excel data to Content Control tags
    for excel_col, tag_name in excel_to_tag_mapping.items():
        value = ""
        
        # Try exact match first
        if excel_col in df_columns:
            value = str(row[excel_col]) if pd.notna(row[excel_col]) else ""
        else:
            # Try similar column names (with space variations)
            excel_col_clean = excel_col.strip().lower()
            found_match = False
            
            # Special handling for location field with extensive matching
            if 'مكان الانعقاد' in excel_col:
                for col in df_columns:
                    # Check for exact match (with or without trailing space)
                    if col == 'مكان الانعقاد ' or col == 'مكان الانعقاد':
                        value = str(row[col]) if pd.notna(row[col]) else ""
                        found_match = True
                        break
                    # Also check if column contains the location keywords
                    elif 'مكان' in str(col) and 'انعقاد' in str(col):
                        value = str(row[col]) if pd.notna(row[col]) else ""
                        found_match = True
                        break
            
            # Special handling for lab field
            elif 'تحتاج لمعمل' in excel_col or 'معمل' in excel_col:
                for col in df_columns:
                    if col == 'تحتاج لمعمل؟' or 'تحتاج لمعمل' in str(col):
                        value = str(row[col]) if pd.notna(row[col]) else ""
                        found_match = True
                        break
            else:
                # Regular matching for other fields
                for col in df_columns:
                    col_clean = str(col).strip().lower()
                    if excel_col_clean == col_clean or excel_col_clean in col_clean or col_clean in excel_col_clean:
                        value = str(row[col]) if pd.notna(row[col]) else ""
                        found_match = True
                        break
        
        # Clean up the value
        if value and str(value).lower() not in ['nan', 'none', 'nat']:
            value = str(value).strip()
        else:
            value = ""
        
        # Special formatting for specific fields
        if value and tag_name == 'مدتها':
            # Add "يوم" after the number for duration field
            try:
                # Try to extract number and add "يوم"
                number = str(value).strip()
                if number.isdigit() or number.replace('.', '').isdigit():
                    value = f"{number} يوم"
                elif number:
                    # If it's not a pure number, just add "يوم" if not already present
                    if 'يوم' not in number:
                        value = f"{number} يوم"
            except:
                # If any error, just use the original value
                pass
        
        # Map to Content Control tag name (only if we have data)
        if value:
            mapping[tag_name] = value
    
    return mapping

def generate_docx_from_template(template_path, mapping, output_name):
    """
    Generate Word document from template using Content Controls
    With complete error recovery and safe table handling
    """
    try:
        # Load template with extreme caution
        try:
            doc = Document(template_path)
        except Exception as load_error:
            st.error(f"❌ خطأ في تحميل القالب: {str(load_error)}")
            return None
        
        replacements_made = 0
        
        # Method 1: Safe Content Controls processing
        try:
            # Use safe iteration to avoid table errors
            content_controls_found = []
            
            # Find all content controls first without processing tables
            try:
                for element in doc.element.body.iter():
                    try:
                        if element.tag.endswith('sdt'):  # Structured Document Tag
                            try:
                                tag_element = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tag')
                                if tag_element is not None:
                                    tag_value = tag_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                    if tag_value:
                                        content_controls_found.append((element, tag_value))
                            except Exception as cc_find_error:
                                # Skip problematic content controls
                                continue
                    except Exception as element_error:
                        # Skip problematic elements
                        continue
                
                # Now process the found content controls
                for element, tag_value in content_controls_found:
                    try:
                        if tag_value.strip() in mapping:
                            data_value = mapping[tag_value.strip()]
                            
                            if data_value:
                                # Find text elements within the content control
                                try:
                                    text_elements = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                    
                                    if text_elements:
                                        # Clear existing text
                                        for text_elem in text_elements:
                                            try:
                                                text_elem.text = ""
                                            except Exception:
                                                continue
                                        
                                        # Set new text (data only)
                                        if len(text_elements) > 0:
                                            try:
                                                text_elements[0].text = data_value
                                                replacements_made += 1
                                            except Exception:
                                                continue
                                except Exception as text_error:
                                    # Skip if text element access fails
                                    continue
                    except Exception as cc_update_error:
                        continue
                        
            except Exception as cc_method_error:
                # Method 2: Safe text replacement fallback (avoid tables completely)
                try:
                    # Only process paragraphs, avoid tables to prevent grid_offset error
                    for paragraph in doc.paragraphs:
                        try:
                            original_text = paragraph.text
                            for tag_name, data_value in mapping.items():
                                try:
                                    if tag_name in paragraph.text and data_value:
                                        paragraph.text = paragraph.text.replace(tag_name, data_value)
                                        replacements_made += 1
                                except Exception:
                                    continue
                        except Exception as para_error:
                            continue
                            
                except Exception as fallback_error:
                    st.error(f"❌ فشل في معالجة المستند: {str(fallback_error)}")
        
        except Exception as processing_error:
            st.error(f"❌ خطأ في معالجة المستند: {str(processing_error)}")
        
        # Show success message only if replacements were made
        if replacements_made > 0:
            st.success(f"✅ تم تحديث {replacements_made} عنصر بنجاح")
        else:
            st.warning("⚠️ لم يتم العثور على عناصر للتحديث")
        
        # Save to BytesIO
        try:
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            return doc_buffer
        except Exception as save_error:
            st.error(f"❌ خطأ في حفظ المستند: {str(save_error)}")
            return None
        
    except Exception as e:
        st.error(f"❌ خطأ عام في توليد المستند: {str(e)}")
        return None

def convert_docx_to_pdf(docx_buffer, output_name):
    """
    Convert DOCX to PDF using docx2pdf
    """
    if not PDF_AVAILABLE:
        return None
    
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_buffer.getvalue())
            temp_docx_path = temp_docx.name
        
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        # Convert to PDF
        convert(temp_docx_path, temp_pdf_path)
        
        # Read PDF content
        with open(temp_pdf_path, 'rb') as pdf_file:
            pdf_content = pdf_file.read()
        
        # Cleanup
        os.unlink(temp_docx_path)
        os.unlink(temp_pdf_path)
        
        return pdf_content
    except Exception as e:
        st.warning(f"فشل في تحويل PDF: {str(e)}")
        return None

# ========================= ENHANCED DASHBOARD FUNCTIONS =========================

def get_status_from_approval_column(status_text):
    """
    Determine course status from "حالة الاعتماد" column values
    """
    if pd.isna(status_text):
        return 'unknown'
    
    status_text = str(status_text).strip().lower()
    
    # Map the specific values from your Excel (handle different spellings)
    if 'مؤكد' in status_text or 'موكد' in status_text:
        return 'confirmed'
    elif 'تاجيل' in status_text or 'تأجيل' in status_text or 'مؤجل' in status_text:
        return 'postponed'
    elif 'تحت الاجراء' in status_text or 'اجراء' in status_text or 'إجراء' in status_text:
        return 'in_progress'
    elif 'ملغ' in status_text or 'الغاء' in status_text or 'إلغاء' in status_text:
        return 'cancelled'
    else:
        return 'unknown'

def get_delivery_method_from_notes(notes_text):
    """
    Determine delivery method from "ملاحظات" column
    Only "عن بعد" is considered remote, everything else is on-site
    """
    if pd.isna(notes_text):
        return 'in_person'
    
    notes_text = str(notes_text).lower().strip()
    
    if 'عن بعد' in notes_text or 'عن بُعد' in notes_text:
        return 'remote'
    else:
        return 'in_person'

def calculate_comprehensive_stats(df, selected_period='all', selected_year=None, selected_month=None, selected_date=None):
    """
    Calculate comprehensive statistics using actual column names from your Excel
    """
    stats = {
        'total_courses': 0,
        'confirmed_courses': 0,
        'postponed_courses': 0,
        'in_progress_courses': 0,
        'cancelled_courses': 0,
        'unknown_courses': 0,
        'remote_courses': 0,
        'in_person_courses': 0,
        'hybrid_courses': 0,
        'current_remote_courses': 0,
        'total_participants': 0,
        'total_training_hours': 0,
        'total_training_days': 0,
        'period_label': ''
    }
    
    if df.empty:
        return stats
    
    # Find relevant columns - exact matching
    start_date_col = None
    approval_status_col = None
    notes_col = None
    participants_col = None
    hours_col = None
    days_col = None
    
    for col in df.columns:
        col_str = str(col).strip()  # Remove trailing spaces
        if col_str == 'تاريخ بداية الدورة بالميلادي':
            start_date_col = col
        elif col_str == 'حالة الاعتماد':
            approval_status_col = col
        elif col_str == 'ملاحظات':
            notes_col = col
        elif 'عدد' in col_str and ('متدرب' in col_str or 'مشارك' in col_str):
            participants_col = col
        elif 'ساعة' in col_str or 'ساعات' in col_str:
            hours_col = col
        elif col_str == 'عدد الايام':
            days_col = col
    
    # Filter data based on selected period
    filtered_df = df.copy()
    
    if start_date_col and selected_period != 'all':
        # Convert dates - handle multiple formats
        def parse_date_flexible(date_str):
            if pd.isna(date_str):
                return pd.NaT
            
            date_str = str(date_str).strip()
            
            # Try different date formats
            for fmt in ['%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y', '%m/%d/%Y']:
                try:
                    return pd.to_datetime(date_str, format=fmt)
                except:
                    continue
            
            # Fallback to pandas auto-parsing
            try:
                return pd.to_datetime(date_str, dayfirst=True)
            except:
                return pd.NaT
        
        filtered_df['parsed_date'] = filtered_df[start_date_col].apply(parse_date_flexible)
        
        if selected_period == 'year' and selected_year:
            filtered_df = filtered_df[filtered_df['parsed_date'].dt.year == selected_year]
            stats['period_label'] = f"سنة {selected_year}"
        elif selected_period == 'month' and selected_year and selected_month:
            filtered_df = filtered_df[
                (filtered_df['parsed_date'].dt.year == selected_year) &
                (filtered_df['parsed_date'].dt.month == selected_month)
            ]
            month_name = calendar.month_name[selected_month]
            stats['period_label'] = f"{month_name} {selected_year}"
        elif selected_period == 'day' and selected_date:
            selected_date = pd.to_datetime(selected_date)
            filtered_df = filtered_df[filtered_df['parsed_date'].dt.date == selected_date.date()]
            stats['period_label'] = f"يوم {selected_date.strftime('%d/%m/%Y')}"
    
    # Calculate basic stats
    stats['total_courses'] = len(filtered_df)
    
    # Status distribution using "حالة الاعتماد"
    if approval_status_col and approval_status_col in filtered_df.columns:
        for _, row in filtered_df.iterrows():
            original_status = row[approval_status_col]
            status = get_status_from_approval_column(original_status)
            
            if status == 'confirmed':
                stats['confirmed_courses'] += 1
            elif status == 'postponed':
                stats['postponed_courses'] += 1
            elif status == 'in_progress':
                stats['in_progress_courses'] += 1
            elif status == 'cancelled':
                stats['cancelled_courses'] += 1
            else:
                stats['unknown_courses'] += 1
    else:
        # If no approval status column found, mark all as unknown
        stats['unknown_courses'] = stats['total_courses']
    
    # Delivery method distribution using "ملاحظات"
    if notes_col and notes_col in filtered_df.columns:
        for _, row in filtered_df.iterrows():
            method = get_delivery_method_from_notes(row[notes_col])
            if method == 'remote':
                stats['remote_courses'] += 1
            elif method == 'in_person':
                stats['in_person_courses'] += 1
            elif method == 'hybrid':
                stats['hybrid_courses'] += 1
    
    # Current remote courses (in_progress + remote)
    if approval_status_col and notes_col:
        for _, row in filtered_df.iterrows():
            status = get_status_from_approval_column(row[approval_status_col])
            method = get_delivery_method_from_notes(row[notes_col])
            if status == 'in_progress' and method == 'remote':
                stats['current_remote_courses'] += 1
    
    # Participants, hours, and days
    if participants_col and participants_col in filtered_df.columns:
        try:
            stats['total_participants'] = int(filtered_df[participants_col].fillna(0).sum())
        except:
            stats['total_participants'] = 0
    
    if hours_col and hours_col in filtered_df.columns:
        try:
            stats['total_training_hours'] = int(filtered_df[hours_col].fillna(0).sum())
        except:
            stats['total_training_hours'] = 0
    
    if days_col and days_col in filtered_df.columns:
        try:
            stats['total_training_days'] = int(filtered_df[days_col].fillna(0).sum())
        except:
            stats['total_training_days'] = 0
    
    return stats

def create_kpi_cards(stats):
    """
    Create KPI cards with enhanced styling using actual data
    """
    # Row 1: Main Course Metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['total_courses']}</div>
            <div class="metric-label">إجمالي الدورات</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['confirmed_courses']}</div>
            <div class="metric-label">دورات مؤكدة</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['postponed_courses']}</div>
            <div class="metric-label">دورات مؤجلة</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['in_progress_courses']}</div>
            <div class="metric-label">تحت الإجراء</div>
        </div>
        """, unsafe_allow_html=True)
    
    # Row 2: Remote Learning Metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['remote_courses']}</div>
            <div class="metric-label">دورات عن بُعد</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['current_remote_courses']}</div>
            <div class="metric-label">دورات عن بُعد حالياً</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['in_person_courses']}</div>
            <div class="metric-label">دورات حضورية</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['total_training_days']}</div>
            <div class="metric-label">إجمالي أيام التدريب</div>
        </div>
        """, unsafe_allow_html=True)

def create_approval_status_distribution_chart(stats):
    """
    Create approval status distribution pie chart using actual data
    """
    status_data = {
        'الحالة': [],
        'العدد': [],
        'اللون': []
    }
    
    # Add data only for non-zero values
    status_mapping = [
        ('مؤكدة', stats['confirmed_courses'], '#2EC4B6'),
        ('مؤجلة', stats['postponed_courses'], '#FFB347'),
        ('تحت الإجراء', stats['in_progress_courses'], '#6A3CBC'),
        ('ملغاة', stats['cancelled_courses'], '#FF6B6B'),
        ('غير محددة', stats['unknown_courses'], '#A8E6CF')
    ]
    
    for label, count, color in status_mapping:
        if count > 0:
            status_data['الحالة'].append(label)
            status_data['العدد'].append(count)
            status_data['اللون'].append(color)
    
    if len(status_data['العدد']) > 0:
        fig = px.pie(
            values=status_data['العدد'],
            names=status_data['الحالة'],
            title="توزيع حالة اعتماد الدورات",
            color_discrete_sequence=status_data['اللون']
        )
        
        fig.update_layout(
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font=dict(color='white', size=12),
            title_font=dict(size=16, color='#6A3CBC'),
            showlegend=True,
            legend=dict(
                orientation="v",
                yanchor="middle",
                y=0.5,
                xanchor="left",
                x=1.01
            )
        )
        
        fig.update_traces(
            textposition='inside',
            textinfo='percent+label',
            hovertemplate='<b>%{label}</b><br>العدد: %{value}<br>النسبة: %{percent}<extra></extra>'
        )
        
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("لا توجد بيانات لعرض توزيع حالة الاعتماد")

def create_delivery_method_chart(stats):
    """
    Create delivery method distribution chart using "ملاحظات" data
    """
    method_data = {
        'طريقة التدريب': [],
        'العدد': [],
        'اللون': []
    }
    
    # Add data only for non-zero values
    method_mapping = [
        ('عن بُعد', stats['remote_courses'], '#6A3CBC'),
        ('حضوري', stats['in_person_courses'], '#2EC4B6'),
        ('مختلط', stats['hybrid_courses'], '#FFB347')
    ]
    
    for label, count, color in method_mapping:
        if count > 0:
            method_data['طريقة التدريب'].append(label)
            method_data['العدد'].append(count)
            method_data['اللون'].append(color)
    
    if len(method_data['العدد']) > 0:
        fig = px.bar(
            x=method_data['العدد'],
            y=method_data['طريقة التدريب'],
            title="توزيع طرق التدريب (من ملاحظات)",
            orientation='h',
            color=method_data['طريقة التدريب'],
            color_discrete_sequence=method_data['اللون']
        )
        
        fig.update_layout(
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font=dict(color='white'),
            title_font=dict(size=16, color='#6A3CBC'),
            showlegend=False,
            xaxis_title="عدد الدورات",
            yaxis_title=""
        )
        
        fig.update_traces(
            hovertemplate='<b>%{y}</b><br>عدد الدورات: %{x}<extra></extra>'
        )
        
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("لا توجد بيانات لعرض توزيع طرق التدريب")

def build_enhanced_dashboard(df):
    """
    Build the enhanced dashboard that properly reads "حالة الاعتماد" data
    """
    st.header("📊 لوحة التحليلات الاحترافية")
    
    # Period selection controls
    st.subheader("🗓️ اختيار الفترة الزمنية")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        period_type = st.selectbox(
            "نوع الفترة",
            ['all', 'year', 'month', 'day'],
            format_func=lambda x: {
                'all': 'جميع الفترات',
                'year': 'سنة محددة',
                'month': 'شهر محدد',
                'day': 'يوم محدد'
            }[x]
        )
    
    selected_year = None
    selected_month = None
    selected_date = None
    
    if period_type in ['year', 'month']:
        with col2:
            # Get available years from data
            if not df.empty and 'تاريخ بداية الدورة بالميلادي' in df.columns:
                df_temp = df.copy()
                df_temp['parsed_date'] = df_temp['تاريخ بداية الدورة بالميلادي'].apply(
                    lambda x: pd.to_datetime(str(x), errors='coerce', dayfirst=True) if pd.notna(x) else pd.NaT
                )
                available_years = sorted(df_temp['parsed_date'].dt.year.dropna().unique())
                
                if available_years:
                    selected_year = st.selectbox("السنة", available_years, index=len(available_years)-1)
    
    if period_type == 'month' and selected_year:
        with col3:
            months = {
                1: 'يناير', 2: 'فبراير', 3: 'مارس', 4: 'أبريل',
                5: 'مايو', 6: 'يونيو', 7: 'يوليو', 8: 'أغسطس',
                9: 'سبتمبر', 10: 'أكتوبر', 11: 'نوفمبر', 12: 'ديسمبر'
            }
            selected_month = st.selectbox(
                "الشهر",
                list(months.keys()),
                format_func=lambda x: months[x],
                index=8  # Default to September
            )
    
    if period_type == 'day':
        with col2:
            selected_date = st.date_input("التاريخ", datetime.now().date())
    
    # Calculate comprehensive statistics
    stats = calculate_comprehensive_stats(
        df, period_type, selected_year, selected_month, selected_date
    )
    
    # Display period label
    if stats['period_label']:
        st.info(f"📅 البيانات المعروضة: {stats['period_label']}")
    
    # KPI Cards
    st.subheader("📈 المؤشرات الرئيسية")
    create_kpi_cards(stats)
    
    # Charts section
    st.subheader("📊 الرسوم البيانية التحليلية")
    
    # First row of charts
    col1, col2 = st.columns(2)
    
    with col1:
        create_approval_status_distribution_chart(stats)
    
    with col2:
        create_delivery_method_chart(stats)
    
    # Summary statistics table
    st.subheader("📋 ملخص تفصيلي")
    
    summary_data = {
        'المؤشر': [
            'إجمالي الدورات',
            'دورات مؤكدة',
            'دورات مؤجلة', 
            'دورات تحت الإجراء',
            'دورات ملغاة',
            'دورات عن بُعد',
            'دورات عن بُعد حالياً',
            'دورات حضورية',
            'إجمالي أيام التدريب'
        ],
        'القيمة': [
            stats['total_courses'],
            stats['confirmed_courses'],
            stats['postponed_courses'],
            stats['in_progress_courses'],
            stats['cancelled_courses'],
            stats['remote_courses'],
            stats['current_remote_courses'],
            stats['in_person_courses'],
            stats['total_training_days']
        ],
        'النسبة المئوية': []
    }
    
    # Calculate percentages
    total = stats['total_courses'] if stats['total_courses'] > 0 else 1
    for i, value in enumerate(summary_data['القيمة']):
        if i < 8:  # For course-related metrics
            percentage = f"{(value / total * 100):.1f}%"
        else:  # For other metrics
            percentage = "-"
        summary_data['النسبة المئوية'].append(percentage)
    
    summary_df = pd.DataFrame(summary_data)
    
    # Display with styling using HTML table to avoid pyarrow dependency
    st.markdown(summary_df.to_html(escape=False, index=False), unsafe_allow_html=True)

# ========================= FORM GENERATOR FUNCTIONS =========================

def build_form_generator(df, template_path):
    """
    Build the accreditation form generator interface
    """
    st.header("📄 مولد نماذج الاعتماد")
    
    if df.empty:
        st.warning("لا توجد بيانات لعرضها")
        return
    
    if not template_path:
        st.warning("يرجى رفع قالب Word أولاً")
        return
    

    # --- FILTERS ---
    filter_col1, filter_col2 = st.columns(2)
    # Filter by audience
    audience_col = None
    for col in df.columns:
        if 'الفئة المستهدفة' in str(col):
            audience_col = col
            break
    if audience_col:
        audience_options = ['الكل'] + sorted([str(x) for x in df[audience_col].dropna().unique()])
        selected_audience = filter_col1.selectbox("فلترة حسب الفئة المستهدفة", audience_options)
        if selected_audience != 'الكل':
            df = df[df[audience_col] == selected_audience]

    # Filter by course start date (day)
    date_col = None
    for col in df.columns:
        if 'تاريخ بداية الدورة بالميلادي' in str(col):
            date_col = col
            break
    selected_day = None
    if date_col:
        # Parse all values to datetime, including string formats like '14/9/2025'
        parsed_dates = pd.to_datetime(df[date_col].astype(str), errors='coerce', dayfirst=True)
        days = parsed_dates.dropna().dt.date.unique()
        day_options = ['الكل'] + [day.strftime('%d/%m/%Y') for day in sorted(days)]
        selected_day = filter_col2.selectbox("فلترة حسب يوم بداية الدورة", day_options)
        if selected_day != 'الكل':
            # Filter by selected day (string format)
            df = df[pd.to_datetime(df[date_col].astype(str), errors='coerce', dayfirst=True).dt.strftime('%d/%m/%Y') == selected_day]

    # Pagination settings
    items_per_page = st.selectbox("عدد العناصر في الصفحة", [5, 10, 20, 50], index=1)

    total_items = len(df)
    total_pages = (total_items - 1) // items_per_page + 1 if total_items > 0 else 1

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        current_page = st.number_input("الصفحة", min_value=1, max_value=total_pages, value=1)

    # Calculate start and end indices
    start_idx = (current_page - 1) * items_per_page
    end_idx = min(start_idx + items_per_page, total_items)

    # Display current page data
    page_df = df.iloc[start_idx:end_idx].copy() if total_items > 0 else pd.DataFrame()
    page_df.reset_index(drop=True, inplace=True)
    
    # Display data with generate buttons
    st.subheader(f"البيانات (الصفحة {current_page} من {total_pages})")
    
    # Extract placeholders from template
    placeholders = extract_placeholders_from_word(template_path)
    
    # Show data table with generate buttons
    for idx, row in page_df.iterrows():
        with st.expander(f"استماره طرح الدوره {start_idx + idx + 1}"):
            col1, col2 = st.columns([3, 1])
            
            with col1:
                # Display row data
                row_dict = row.to_dict()
                for key, value in row_dict.items():
                    st.text(f"{key}: {value}")
            
            with col2:
                if st.button(f"⬇️ توليد النموذج", key=f"generate_{start_idx + idx}"):
                    # Generate mapping
                    mapping = build_mapping(row, df.columns.tolist())
                    
                    # Generate DOCX
                    output_name = f"استماره_طرح_الدوره_{start_idx + idx + 1}"
                    docx_buffer = generate_docx_from_template(template_path, mapping, output_name)
                    
                    if docx_buffer:
                        # Download DOCX
                        st.download_button(
                            label="📄 تحميل Word",
                            data=docx_buffer.getvalue(),
                            file_name=f"{output_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_docx_{start_idx + idx}"
                        )
                        
                        # Generate PDF if available
                        if PDF_AVAILABLE:
                            pdf_content = convert_docx_to_pdf(docx_buffer, output_name)
                            if pdf_content:
                                st.download_button(
                                    label="📄 تحميل PDF",
                                    data=pdf_content,
                                    file_name=f"{output_name}.pdf",
                                    mime="application/pdf",
                                    key=f"download_pdf_{start_idx + idx}"
                                )
    
    # Bulk generation option
    st.subheader("📦 التوليد المجمع")
    if st.button("توليد جميع النماذج وتحميل ملف ZIP"):
        with st.spinner("جاري توليد النماذج..."):
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for idx, row in df.iterrows():
                    try:
                        # Generate mapping
                        mapping = build_mapping(row, df.columns.tolist())
                        
                        # Generate DOCX
                        output_name = f"استماره_طرح_الدوره_{idx + 1}"
                        docx_buffer = generate_docx_from_template(template_path, mapping, output_name)
                        
                        if docx_buffer:
                            zip_file.writestr(f"{output_name}.docx", docx_buffer.getvalue())
                            
                            # Add PDF if available
                            if PDF_AVAILABLE:
                                pdf_content = convert_docx_to_pdf(docx_buffer, output_name)
                                if pdf_content:
                                    zip_file.writestr(f"{output_name}.pdf", pdf_content)
                    except Exception as e:
                        st.warning(f"خطأ في توليد النموذج للسطر {idx + 1}: {str(e)}")
            
            zip_buffer.seek(0)
            
            st.download_button(
                label="📦 تحميل جميع النماذج (ZIP)",
                data=zip_buffer.getvalue(),
                file_name="جميع_النماذج.zip",
                mime="application/zip"
            )

# ========================= COMPARISON FUNCTIONS =========================

def build_comparison_view(df, template_path):
    """
    Build Excel vs Word comparison interface
    """
    st.header("🔍 مقارنة Excel مع Word")
    
    if not template_path:
        st.warning("يرجى رفع قالب Word أولاً")
        return
    
    # Extract placeholders from Word template
    placeholders = extract_placeholders_from_word(template_path)
    
    if not placeholders:
        st.warning("لم يتم العثور على أي عناصر نائبة في قالب Word")
        return
    
    # Get Excel columns
    excel_columns = df.columns.tolist() if not df.empty else []
    
    # Build comparison data
    comparison_data = []
    matched_columns = set()
    
    for placeholder in placeholders:
        # Try to find exact match
        if placeholder in excel_columns:
            comparison_data.append({
                "العنصر النائب في Word": f"{{{{{placeholder}}}}}",
                "عمود Excel المطابق": placeholder,
                "الحالة": "✅ مطابق"
            })
            matched_columns.add(placeholder)
        else:
            # Try to find similar column
            similar_cols = [col for col in excel_columns if placeholder.lower() in str(col).lower() or str(col).lower() in placeholder.lower()]
            if similar_cols:
                comparison_data.append({
                    "العنصر النائب في Word": f"{{{{{placeholder}}}}}",
                    "عمود Excel المطابق": similar_cols[0],
                    "الحالة": "⚠️ مشابه"
                })
                matched_columns.add(similar_cols[0])
            else:
                comparison_data.append({
                    "العنصر النائب في Word": f"{{{{{placeholder}}}}}",
                    "عمود Excel المطابق": "غير موجود",
                    "الحالة": "❌ مفقود"
                })
    
    # Add unused Excel columns
    unused_columns = set(excel_columns) - matched_columns
    for col in unused_columns:
        comparison_data.append({
            "العنصر النائب في Word": "غير موجود",
            "عمود Excel المطابق": col,
            "الحالة": "⚠️ غير مستخدم"
        })
    
    # Display comparison table
    st.subheader("📊 جدول المقارنة")
    comparison_df = pd.DataFrame(comparison_data)
    
    # Apply styling based on status
    def style_status(val):
        if "✅" in val:
            return 'color: #28a745; font-weight: bold'
        elif "❌" in val:
            return 'color: #dc3545; font-weight: bold'
        elif "⚠️" in val:
            return 'color: #ffc107; font-weight: bold'
        return ''
    
    # Display the comparison table using HTML to avoid pyarrow dependency
    st.write(comparison_df.to_html(escape=False), unsafe_allow_html=True)
    
    # Summary statistics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        matched_count = len([item for item in comparison_data if "✅" in item["الحالة"]])
        st.metric("عناصر مطابقة", matched_count)
    
    with col2:
        similar_count = len([item for item in comparison_data if "⚠️" in item["الحالة"] and "مشابه" in item["الحالة"]])
        st.metric("عناصر مشابهة", similar_count)
    
    with col3:
        missing_count = len([item for item in comparison_data if "❌" in item["الحالة"]])
        st.metric("عناصر مفقودة", missing_count)
    
    with col4:
        unused_count = len([item for item in comparison_data if "غير مستخدم" in item["الحالة"]])
        st.metric("أعمدة غير مستخدمة", unused_count)
    
    # Warnings and recommendations
    if missing_count > 0:
        st.warning(f"⚠️ يوجد {missing_count} عنصر نائب في قالب Word لا يوجد له عمود مطابق في Excel")
    
    if unused_count > 0:
        st.info(f"💡 يوجد {unused_count} عمود في Excel لا يستخدم في قالب Word")
    
    # Export comparison report
    if st.button("📄 تصدير تقرير المقارنة"):
        excel_buffer = io.BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
            comparison_df.to_excel(writer, sheet_name='تقرير المقارنة', index=False)
        
        excel_buffer.seek(0)
        st.download_button(
            label="📊 تحميل تقرير Excel",
            data=excel_buffer.getvalue(),
            file_name="تقرير_المقارنة.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

# ========================= MAIN APPLICATION =========================

def main():
    """
    Main application function
    """
    # Beautiful header with logo and gradient background
    logo_path = "assets/logo.png"
    
    # Check if logo exists
    if os.path.exists(logo_path):
        # Create header with logo
        st.markdown(f"""
        <div class="header-container">
            <div class="logo-container">
                <img src="data:image/png;base64,{get_base64_of_image(logo_path)}" class="logo-image">
            </div>
            <h1 class="header-title">نظام إدارة الدورات التدريبية</h1>
            <p class="header-subtitle">Training Courses Management System</p>
        </div>
        """, unsafe_allow_html=True)
    else:
        # Fallback header without logo
        st.markdown("""
        <div class="header-container">
            <h1 class="header-title">📚 نظام إدارة الدورات التدريبية</h1>
            <p class="header-subtitle">Training Courses Management System</p>
        </div>
        """, unsafe_allow_html=True)
    
    # Sidebar for file uploads and settings
    with st.sidebar:
        # Enhanced sidebar header
        st.markdown("""
        <div style="text-align: center; padding: 1rem; background: linear-gradient(135deg, #6A3CBC 0%, #2EC4B6 100%); border-radius: 10px; margin-bottom: 1rem;">
            <h2 style="color: white; margin: 0; font-size: 1.5rem;">⚙️ الإعدادات</h2>
            <p style="color: rgba(255,255,255,0.8); margin: 0.5rem 0 0 0; font-size: 0.9rem;">إعدادات النظام والملفات</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Check for default files first with complete error handling
        try:
            if hasattr(config, 'EXCEL_FILE_PATH') and hasattr(config, 'TEMPLATE_FILE_PATH'):
                default_excel_path = config.EXCEL_FILE_PATH
                default_template_path = config.TEMPLATE_FILE_PATH
                
                # Test template loading early to catch errors safely
                try:
                    # Just test if we can open the document without processing its content
                    test_doc = Document(default_template_path)
                    # Test basic access without table iteration
                    _ = len(test_doc.paragraphs)
                    st.success("✅ قالب Word تم تحميله بنجاح")
                    # Don't test tables here to avoid grid_offset errors
                except Exception as template_test_error:
                    st.warning(f"⚠️ تحذير من قالب Word: {str(template_test_error)}")
                    st.info("🔄 سيتم تجربة القالب مع التعامل الآمن مع الأخطاء")
                    # Don't set to None - still try to use it with safe processing
            else:
                st.error("❌ خطأ في إعدادات الملفات")
                default_excel_path = None
                default_template_path = None
        except Exception as config_error:
            st.error(f"❌ خطأ في قراءة الإعدادات: {str(config_error)}")
            default_excel_path = None
            default_template_path = None
        
        excel_df = pd.DataFrame()
        available_sheets = []
        excel_path = None
        template_path = None
        
        # Try to load default Excel file if it exists
        if default_excel_path and os.path.exists(default_excel_path):
            st.info(f"📊 تم العثور على ملف البيانات: {os.path.basename(default_excel_path)}")
            excel_path = default_excel_path
            available_sheets = get_available_sheets(excel_path)
            
            if available_sheets:
                # Auto-select September sheet if available, otherwise show selector
                september_sheet = None
                september_variations = ["سبتمبر", "September", "9", "09"]
                
                for variation in september_variations:
                    if variation in available_sheets:
                        september_sheet = variation
                        break
                
                if september_sheet:
                    st.success(f"🗓️ تم تحديد شهر سبتمبر تلقائياً: '{september_sheet}'")
                    selected_sheet = september_sheet
                    # Also show option to change if needed
                    if st.checkbox("تغيير الشهر", key="change_month"):
                        selected_sheet = st.selectbox("اختر الشهر (ورقة العمل)", available_sheets, 
                                                    index=available_sheets.index(september_sheet))
                else:
                    st.warning("لم يتم العثور على شهر سبتمبر، يرجى اختيار الشهر:")
                    selected_sheet = st.selectbox("اختر الشهر (ورقة العمل)", available_sheets)
                
                excel_df = load_excel_data(excel_path, selected_sheet)
                
                if not excel_df.empty:
                    st.success(f"تم تحميل {len(excel_df)} صف من بيانات شهر: {selected_sheet}")
                    
                    # Show data summary
                    st.info(f"📋 ملخص البيانات: {len(excel_df)} دورة في {selected_sheet}")
        else:
            st.warning("لم يتم العثور على ملف البيانات الافتراضي")
        
        # Excel file upload (alternative)
        st.markdown("**أو قم برفع ملف Excel آخر:**")
        excel_file = st.file_uploader("رفع ملف Excel", type=['xlsx', 'xls'])
        
        if excel_file:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
                tmp.write(excel_file.getvalue())
                excel_path = tmp.name
            
            # Get available sheets
            available_sheets = get_available_sheets(excel_path)
            
            if available_sheets:
                # Auto-select September sheet if available
                september_sheet = None
                september_variations = ["سبتمبر", "September", "9", "09"]
                
                for variation in september_variations:
                    if variation in available_sheets:
                        september_sheet = variation
                        break
                
                if september_sheet:
                    st.success(f"🗓️ تم تحديد شهر سبتمبر تلقائياً: '{september_sheet}'")
                    default_index = available_sheets.index(september_sheet)
                else:
                    default_index = 0
                
                selected_sheet = st.selectbox("اختر الشهر (ورقة العمل)", available_sheets, 
                                            index=default_index, key="uploaded_sheet")
                excel_df = load_excel_data(excel_path, selected_sheet)
                
                if not excel_df.empty:
                    st.success(f"تم تحميل {len(excel_df)} صف من البيانات")
                    
                    # Show preview
                    st.subheader("📋 معاينة البيانات")
                    st.write(f"عدد الصفوف: {len(excel_df)}")
                    st.write(f"عدد الأعمدة: {len(excel_df.columns)}")
                    
                    # Display first few rows using HTML table to avoid pyarrow
                    html_table = excel_df.head().to_html(escape=False, index=False)
                    st.markdown("**أول 5 صفوف:**")
                    st.markdown(html_table, unsafe_allow_html=True)
            
            # Clean up temp file
            try:
                os.unlink(excel_path)
            except:
                pass
        
        # Check for default Word template
        if default_template_path and os.path.exists(default_template_path):
            st.info(f"📄 تم العثور على قالب Word: {os.path.basename(default_template_path)}")
            template_path = default_template_path
        
        # Word template upload (alternative)
        st.markdown("**أو قم برفع قالب Word آخر:**")
        template_file = st.file_uploader("رفع قالب Word", type=['docx'])
        
        if template_file:
            # Save uploaded template temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(template_file.getvalue())
                template_path = tmp.name
            
            st.success("تم رفع قالب Word بنجاح")
        
        # Export options
        st.markdown("---")
        st.subheader("📤 خيارات التصدير")
        
        if not excel_df.empty:
            # Monthly summary export
            if st.button("تصدير الملخص الشهري"):
                monthly_stats = calculate_monthly_stats(excel_df)
                
                summary_data = {
                    'المؤشر': ['إجمالي الدورات المخططة', 'دورات منفذة', 'دورات ملغاة', 'دورات مؤجلة', 'إجمالي أيام التدريب'],
                    'القيمة': [monthly_stats['total_planned'], monthly_stats['executed'], 
                              monthly_stats['cancelled'], monthly_stats['postponed'], monthly_stats['total_training_days']]
                }
                
                summary_df = pd.DataFrame(summary_data)
                
                excel_buffer = io.BytesIO()
                with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
                    summary_df.to_excel(writer, sheet_name='الملخص الشهري', index=False)
                    excel_df.to_excel(writer, sheet_name='البيانات الكاملة', index=False)
                
                excel_buffer.seek(0)
                st.download_button(
                    label="📊 تحميل الملخص الشهري",
                    data=excel_buffer.getvalue(),
                    file_name=f"الملخص_الشهري_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
    
    # Main content tabs
    tab1, tab2 = st.tabs(["📊 لوحة الإحصائيات", "📄 توليد النماذج"])
    
    with tab1:
        build_enhanced_dashboard(excel_df)
    
    with tab2:
        build_form_generator(excel_df, template_path)
    
    # Cleanup temporary files (only if they were uploaded, not default files)
    if excel_file and excel_path and excel_path != config.EXCEL_FILE_PATH:
        try:
            os.unlink(excel_path)
        except:
            pass
    
    if template_file and template_path and template_path != config.TEMPLATE_FILE_PATH:
        try:
            os.unlink(template_path)
        except:
            pass

if __name__ == "__main__":
    main()
