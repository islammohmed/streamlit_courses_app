from docx import Document
from docx.shared import Inches

def create_sample_template():
    """Create a sample Word template with Arabic placeholders"""
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('نموذج شهادة إتمام الدورة التدريبية', 0)
    title.alignment = 2  # Right alignment
    
    # Add content with placeholders
    doc.add_paragraph('')
    
    # Header information
    doc.add_paragraph('وزارة التعليم - معهد التدريب المتقدم')
    doc.add_paragraph('إدارة التطوير والتدريب')
    doc.add_paragraph('')
    
    # Certificate content
    doc.add_paragraph('شهادة إتمام', style='Heading 1')
    doc.add_paragraph('')
    
    content = """
نشهد أن المتدرب/ة: {{اسم_المتدرب}}
قد أتم/ت بنجاح الدورة التدريبية:

اسم البرنامج: {{اسم_البرنامج}}
كود الدورة: {{كود_الدورة}}
المدرب: {{المدرب}}
الجمهور المستهدف: {{الجمهور_المستهدف}}

تفاصيل الدورة:
- تاريخ البداية: {{تاريخ_البداية}}
- تاريخ النهاية: {{تاريخ_النهاية}}
- عدد الساعات التدريبية: {{عدد_الساعات}} ساعة
- مكان التدريب: {{المكان}}
- عدد المتدربين: {{عدد_المتدربين}} متدرب/ة

ملاحظات إضافية: {{ملاحظات}}

نتمنى للمتدرب/ة التوفيق في تطبيق ما تعلمه في العمل.
"""
    
    doc.add_paragraph(content)
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Footer
    footer_table = doc.add_table(rows=1, cols=3)
    footer_table.cell(0, 0).text = 'مدير المعهد\n_________________'
    footer_table.cell(0, 1).text = 'مدير التدريب\n_________________'
    footer_table.cell(0, 2).text = 'التاريخ: {{تاريخ_الاصدار}}\n_________________'
    
    # Save the document
    doc.save('sample_template.docx')
    print("Sample Word template created: sample_template.docx")
    print("Placeholders included:")
    placeholders = [
        'اسم_المتدرب', 'اسم_البرنامج', 'كود_الدورة', 'المدرب', 
        'الجمهور_المستهدف', 'تاريخ_البداية', 'تاريخ_النهاية', 
        'عدد_الساعات', 'المكان', 'عدد_المتدربين', 'ملاحظات', 'تاريخ_الاصدار'
    ]
    for placeholder in placeholders:
        print(f"  - {{{{{placeholder}}}}}")

if __name__ == "__main__":
    create_sample_template()
